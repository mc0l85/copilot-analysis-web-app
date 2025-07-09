I'd like to convert a python script with GUI to a PHP front-end that uses a python script in the back end.  I'll use it on an ubuntu lxc container.

I don't want the web server to save any data.  The input data files must be deleted after the processing is complete.  The excel file and leaderboard file should be deleted after 5 minutes.  The php page should provide links to both files once they are available.

Python:
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
from docx.shared import Inches
from datetime import datetime, timedelta
import os
import threading
import subprocess
import sys
import webbrowser
from openpyxl.utils import get_column_letter
from openpyxl.styles import PatternFill, Font, Alignment
from openpyxl.formatting.rule import ColorScaleRule, DataBarRule

class CopilotAnalysisApp:
    """
    A GUI application for analyzing Microsoft Copilot usage data.
    """
    def __init__(self, master):
        """
        Initializes the application window and widgets.
        """
        self.master = master
        master.title("Copilot Usage Analysis Tool")
        master.geometry("950x900") # Increased height for new filters

        self.notebook = ttk.Notebook(master, padding=10)
        self.notebook.pack(expand=True, fill="both")
        
        self.analysis_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(self.analysis_frame, text='Analysis')
        
        self.leaderboard_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(self.leaderboard_frame, text='Leaderboard')

        self.deep_dive_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(self.deep_dive_frame, text='Deep Dive')

        self.setup_analysis_tab()
        self.setup_leaderboard_tab()
        self.setup_deep_dive_tab()
        
        self.target_users_file = ""
        self.usage_report_files = []
        self.full_usage_data = None
        self.target_user_data = None
        self.utilized_metrics_df = None
        self.output_folder_path = os.path.join(os.getcwd(), "Copilot_Analysis_Outputs")
        self.leaderboard_html_path = ""
        
        # Initialize lists for all filters
        self.all_companies = []
        self.all_departments = []
        self.all_locations = []
        self.all_managers = []

    def setup_analysis_tab(self):
        file_selection_frame = tk.LabelFrame(self.analysis_frame, text="File Selection", padx=10, pady=10)
        file_selection_frame.pack(padx=10, pady=10, fill="x")

        target_users_label = tk.Label(file_selection_frame, text="Target Users File (Optional):")
        target_users_label.grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.target_users_path = tk.Entry(file_selection_frame, width=80)
        self.target_users_path.grid(row=0, column=1, columnspan=2, padx=5, pady=5, sticky="ew")
        target_users_button = tk.Button(file_selection_frame, text="Browse", command=self.load_target_users)
        target_users_button.grid(row=0, column=3, padx=5, pady=5)
        
        # Main container for all filter listboxes
        filter_frame = tk.Frame(file_selection_frame)
        filter_frame.grid(row=1, column=0, columnspan=4, sticky='ew', pady=10)
        filter_frame.columnconfigure(0, weight=1)
        filter_frame.columnconfigure(1, weight=1)
        
        # --- ROW 0 of Filters ---
        company_frame = tk.LabelFrame(filter_frame, text="Company Filter (Optional, Ctrl+Click for multi-select)")
        company_frame.grid(row=0, column=0, padx=5, pady=5, sticky='nsew')
        company_scrollbar = tk.Scrollbar(company_frame)
        company_scrollbar.pack(side="right", fill="y")
        self.company_listbox = tk.Listbox(company_frame, selectmode=tk.EXTENDED, yscrollcommand=company_scrollbar.set, height=5)
        self.company_listbox.pack(side="left", fill="both", expand=True)
        company_scrollbar.config(command=self.company_listbox.yview)

        department_frame = tk.LabelFrame(filter_frame, text="Department Filter (Optional, Ctrl+Click for multi-select)")
        department_frame.grid(row=0, column=1, padx=5, pady=5, sticky='nsew')
        department_scrollbar = tk.Scrollbar(department_frame)
        department_scrollbar.pack(side="right", fill="y")
        self.department_listbox = tk.Listbox(department_frame, selectmode=tk.EXTENDED, yscrollcommand=department_scrollbar.set, height=5)
        self.department_listbox.pack(side="left", fill="both", expand=True)
        department_scrollbar.config(command=self.department_listbox.yview)

        # --- ROW 1 of Filters ---
        location_frame = tk.LabelFrame(filter_frame, text="Location Filter (Optional, Ctrl+Click for multi-select)")
        location_frame.grid(row=1, column=0, padx=5, pady=5, sticky='nsew')
        location_scrollbar = tk.Scrollbar(location_frame)
        location_scrollbar.pack(side="right", fill="y")
        self.location_listbox = tk.Listbox(location_frame, selectmode=tk.EXTENDED, yscrollcommand=location_scrollbar.set, height=5)
        self.location_listbox.pack(side="left", fill="both", expand=True)
        location_scrollbar.config(command=self.location_listbox.yview)

        manager_frame = tk.LabelFrame(filter_frame, text="Manager Filter (Optional, Ctrl+Click for multi-select)")
        manager_frame.grid(row=1, column=1, padx=5, pady=5, sticky='nsew')
        manager_scrollbar = tk.Scrollbar(manager_frame)
        manager_scrollbar.pack(side="right", fill="y")
        self.manager_listbox = tk.Listbox(manager_frame, selectmode=tk.EXTENDED, yscrollcommand=manager_scrollbar.set, height=5)
        self.manager_listbox.pack(side="left", fill="both", expand=True)
        manager_scrollbar.config(command=self.manager_listbox.yview)

        usage_reports_label = tk.Label(file_selection_frame, text="Usage Reports (CSV or Excel):")
        usage_reports_label.grid(row=2, column=0, sticky="w", padx=5, pady=5)
        self.usage_reports_path = tk.Entry(file_selection_frame, width=80)
        self.usage_reports_path.grid(row=2, column=1, columnspan=2, padx=5, pady=5, sticky="ew")
        usage_reports_button = tk.Button(file_selection_frame, text="Browse", command=self.load_usage_reports)
        usage_reports_button.grid(row=2, column=3, padx=5, pady=5)

        file_selection_frame.columnconfigure(1, weight=1)
        
        control_frame = tk.Frame(self.analysis_frame, padx=10, pady=10)
        control_frame.pack(padx=10, pady=5, fill="x")
        
        self.run_button = tk.Button(control_frame, text="Run Analysis", command=self.run_analysis_thread, font=("Helvetica", 12, "bold"))
        self.run_button.pack(side="left", padx=10)

        self.open_folder_button = tk.Button(control_frame, text="Open Report Folder", command=self.open_report_folder, font=("Helvetica", 12))
        
        log_frame = tk.LabelFrame(self.analysis_frame, text="Log", padx=10, pady=10)
        log_frame.pack(padx=10, pady=10, fill="both", expand=True)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, state='disabled', wrap=tk.WORD, height=15)
        self.log_text.pack(fill="both", expand=True)

    def setup_leaderboard_tab(self):
        info_label = tk.Label(self.leaderboard_frame, text="The leaderboard is generated as a shareable HTML file.\nRun an analysis to activate the buttons below.", justify="left", font=("Helvetica", 10))
        info_label.pack(pady=10, anchor="w")
        
        button_frame = tk.Frame(self.leaderboard_frame)
        button_frame.pack(pady=20)

        self.open_leaderboard_button = tk.Button(button_frame, text="Open Leaderboard in Browser", command=self.open_leaderboard_file, state='disabled', font=("Helvetica", 12, "bold"))
        self.open_leaderboard_button.pack(pady=10)

        self.copy_html_button = tk.Button(button_frame, text="Copy HTML Source to Clipboard", command=self.copy_leaderboard_html, state='disabled', font=("Helvetica", 12))
        self.copy_html_button.pack(pady=10)

    def setup_deep_dive_tab(self):
        search_frame_top = tk.Frame(self.deep_dive_frame)
        search_frame_top.pack(fill="x", pady=5)

        deep_dive_info = tk.Label(search_frame_top, text="Enter a user's email for a detailed analysis.\nRun an analysis on the 'Analysis' tab first to load data.", justify="left")
        deep_dive_info.pack(pady=5, anchor="w")

        search_frame = tk.Frame(search_frame_top)
        search_frame.pack(fill="x", pady=5)
        
        user_email_label = tk.Label(search_frame, text="User Email:")
        user_email_label.pack(side="left", padx=5)
        
        self.user_email_entry = tk.Entry(search_frame, width=50)
        self.user_email_entry.pack(side="left", expand=True, fill="x", padx=5)
        
        search_button = tk.Button(search_frame, text="Search", command=self.perform_deep_dive)
        search_button.pack(side="left", padx=5)
        
        paned_window = ttk.PanedWindow(self.deep_dive_frame, orient=tk.VERTICAL)
        paned_window.pack(fill="both", expand=True, pady=10)

        graph_frame = tk.LabelFrame(paned_window, text="Usage Complexity Trend", padx=10, pady=10)
        paned_window.add(graph_frame, weight=1)

        results_frame = tk.LabelFrame(paned_window, text="Statistics & Details", padx=10, pady=10)
        paned_window.add(results_frame, weight=1)
        
        self.fig = Figure(figsize=(8, 4), dpi=100)
        self.ax = self.fig.add_subplot(111)
        self.canvas = FigureCanvasTkAgg(self.fig, master=graph_frame)
        self.canvas.get_tk_widget().pack(fill="both", expand=True)
        
        self.deep_dive_results = scrolledtext.ScrolledText(results_frame, state='disabled', wrap=tk.WORD, height=15)
        self.deep_dive_results.pack(fill="both", expand=True)

    def log(self, message):
        self.master.after(0, self._log_message, message)

    def _log_message(self, message):
        self.log_text.config(state='normal')
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.config(state='disabled')
        self.log_text.see(tk.END)

    def load_target_users(self):
        filepath = filedialog.askopenfilename(
            title="Select Target Users File",
            filetypes=(("CSV and Text files", "*.csv *.txt"), ("All files", "*.*"))
        )
        if not filepath:
            return

        self.target_users_file = filepath
        self.target_users_path.delete(0, tk.END)
        self.target_users_path.insert(0, filepath)
        
        try:
            self.target_user_data = pd.read_csv(filepath)
            
            # Update required columns to include Company and Department
            required_cols = ['UserPrincipalName', 'Company', 'Department', 'City', 'ManagerLine']
            if all(col in self.target_user_data.columns for col in required_cols):
                # Populate Company filter
                self.company_listbox.delete(0, tk.END)
                self.all_companies = sorted(self.target_user_data['Company'].dropna().unique().tolist())
                for company in self.all_companies:
                    self.company_listbox.insert(tk.END, company)

                # Populate Department filter
                self.department_listbox.delete(0, tk.END)
                self.all_departments = sorted(self.target_user_data['Department'].dropna().unique().tolist())
                for department in self.all_departments:
                    self.department_listbox.insert(tk.END, department)

                # Populate Location filter
                self.location_listbox.delete(0, tk.END)
                self.all_locations = sorted(self.target_user_data['City'].dropna().unique().tolist())
                for location in self.all_locations:
                    self.location_listbox.insert(tk.END, location)
                
                # Populate Manager filter
                self.manager_listbox.delete(0, tk.END)
                all_managers = set()
                for chain in self.target_user_data['ManagerLine'].dropna():
                    managers_in_chain = [m.strip() for m in chain.split('->')]
                    all_managers.update(managers_in_chain)
                self.all_managers = sorted(list(all_managers))
                for manager in self.all_managers:
                    self.manager_listbox.insert(tk.END, manager)
                
                self.log("Target user file loaded. Filters populated.")
            else:
                missing_cols = [col for col in required_cols if col not in self.target_user_data.columns]
                self.log(f"Target user file missing required columns: {missing_cols}")
                messagebox.showerror("Column Error", f"Target user file must contain the columns: {', '.join(required_cols)}")
                self.target_user_data = None
        except Exception as e:
            self.log(f"Error reading target user file: {e}")
            messagebox.showerror("File Read Error", f"Could not read the target user file.\n\nError: {e}")

    def load_usage_reports(self):
        filepaths = filedialog.askopenfilenames(
            title="Select Usage Reports",
            filetypes=(("Report Files", "*.csv *.xlsx *.xls"), ("All files", "*.*"))
        )
        if filepaths:
            self.usage_report_files = filepaths
            self.usage_reports_path.delete(0, tk.END)
            self.usage_reports_path.insert(0, "; ".join(filepaths))
            self.log(f"Selected {len(filepaths)} usage report(s).")

    def run_analysis_thread(self):
        if not self.usage_report_files:
            messagebox.showerror("Error", "Please select one or more Usage Report files.")
            return
        self.run_button.config(state="disabled", text="Running...")
        self.log_text.config(state='normal')
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state='disabled')
        analysis_thread = threading.Thread(target=self.execute_analysis)
        analysis_thread.start()

    def execute_analysis(self):
        try:
            self.log("--- Starting Analysis ---")
            
            self.log("1. Loading data...")
            all_reports = []
            for file in self.usage_report_files:
                try:
                    if file.lower().endswith('.csv'):
                        df = pd.read_csv(file)
                    else:
                        df = pd.read_excel(file)
                    all_reports.append(df)
                except Exception as e:
                    self.log(f"Could not read file: {os.path.basename(file)}. Error: {e}")
                    continue
            
            if not all_reports:
                messagebox.showerror("Error", "No usage reports could be read. Please check the file formats.")
                self.reset_button()
                return

            usage_df = pd.concat(all_reports, ignore_index=True)
            usage_df['User Principal Name'] = usage_df['User Principal Name'].str.lower()
            
            self.log("2. Preparing data...")
            date_cols = [col for col in usage_df.columns if 'date' in col.lower()]
            for col in date_cols:
                usage_df[col] = pd.to_datetime(usage_df[col], errors='coerce', format='mixed')
            
            self.full_usage_data = usage_df.copy()
            self.log("Data loaded successfully.")

            if self.target_user_data is not None:
                self.log("Target user file provided. Filtering user list...")
                filtered_target_df = self.target_user_data.copy()
                
                # Apply Company Filter
                selected_company_indices = self.company_listbox.curselection()
                if selected_company_indices:
                    selected_companies = [self.all_companies[i] for i in selected_company_indices]
                    self.log(f"Filtering for companies: {', '.join(selected_companies)}")
                    filtered_target_df = filtered_target_df[filtered_target_df['Company'].isin(selected_companies)]

                # Apply Department Filter
                selected_department_indices = self.department_listbox.curselection()
                if selected_department_indices:
                    selected_departments = [self.all_departments[i] for i in selected_department_indices]
                    self.log(f"Filtering for departments: {', '.join(selected_departments)}")
                    filtered_target_df = filtered_target_df[filtered_target_df['Department'].isin(selected_departments)]

                # Apply Location Filter
                selected_location_indices = self.location_listbox.curselection()
                if selected_location_indices:
                    selected_locations = [self.all_locations[i] for i in selected_location_indices]
                    self.log(f"Filtering for locations: {', '.join(selected_locations)}")
                    filtered_target_df = filtered_target_df[filtered_target_df['City'].isin(selected_locations)]
                
                # Apply Manager Filter
                selected_manager_indices = self.manager_listbox.curselection()
                if selected_manager_indices:
                    selected_managers = [self.all_managers[i] for i in selected_manager_indices]
                    self.log(f"Filtering for managers: {', '.join(selected_managers)}")
                    
                    manager_mask = pd.Series([False] * len(filtered_target_df), index=filtered_target_df.index)
                    for manager in selected_managers:
                        manager_mask |= filtered_target_df['ManagerLine'].str.contains(f'\\b{manager}\\b', regex=True, na=False)
                    filtered_target_df = filtered_target_df[manager_mask]

                target_users_emails = set(filtered_target_df['UserPrincipalName'].str.lower())
                all_report_emails = set(usage_df['User Principal Name'].unique())
                utilized_emails = target_users_emails.intersection(all_report_emails)
                self.log(f"Analyzing {len(utilized_emails)} of {len(target_users_emails)} target users found in reports.")
            else:
                self.log("No target user file provided. Analyzing all users from reports.")
                utilized_emails = set(usage_df['User Principal Name'].unique())

            if not utilized_emails:
                self.log("Error: No matching users found to analyze. Please check your input files and filters.")
                messagebox.showerror("No Data", "No users were found to analyze based on your selections.")
                self.reset_button()
                return

            matched_users_df = usage_df[usage_df['User Principal Name'].isin(utilized_emails)].copy()
            copilot_tool_cols = [col for col in matched_users_df.columns if 'Last activity date of' in col]

            min_report_date = usage_df['Report Refresh Date'].min()
            max_report_date = usage_df['Report Refresh Date'].max()
            total_months_in_period = (max_report_date.year - min_report_date.year) * 12 + max_report_date.month - min_report_date.month + 1

            user_metrics = []
            for email in utilized_emails:
                user_data = matched_users_df[matched_users_df['User Principal Name'] == email]
                
                activity_dates = user_data[copilot_tool_cols].stack().dropna().unique()
                activity_dates = pd.to_datetime(activity_dates)

                if len(activity_dates) == 0:
                    first_activity, last_activity = pd.NaT, pd.NaT
                    active_months, complexity, avg_tools_per_month, trend = 0, 0, 0, "N/A"
                    report_dates = sorted(user_data['Report Refresh Date'].unique())
                    first_activity = report_dates[0] if report_dates else pd.NaT
                else:
                    first_activity, last_activity = activity_dates.min(), activity_dates.max()
                    active_months = len(pd.to_datetime(activity_dates).to_period('M').unique())
                    complexity = user_data[copilot_tool_cols].notna().any().sum()
                    
                    monthly_activity = user_data[copilot_tool_cols].stack().dropna().reset_index().rename(columns={'level_1': 'Tool', 0: 'Date'})
                    monthly_activity['Month'] = pd.to_datetime(monthly_activity['Date']).dt.to_period('M')
                    avg_tools_per_month = monthly_activity.groupby('Month')['Tool'].nunique().mean() if not monthly_activity.empty else 0

                    trend = "N/A"
                    if len(activity_dates) > 1:
                        trend = "Stable"
                        timeline_midpoint = first_activity + (last_activity - first_activity) / 2
                        first_half_activity = monthly_activity[monthly_activity['Date'] <= timeline_midpoint]
                        second_half_activity = monthly_activity[monthly_activity['Date'] > timeline_midpoint]
                        if second_half_activity['Tool'].nunique() > first_half_activity['Tool'].nunique():
                            trend = "Increasing"
                        elif second_half_activity['Tool'].nunique() < first_half_activity['Tool'].nunique():
                            trend = "Decreasing"

                consistency = (active_months / total_months_in_period) * 100 if total_months_in_period > 0 else 0
                
                user_metrics.append({
                    'Email': email, 
                    'Usage Consistency (%)': consistency, 
                    'Overall Recency': last_activity, 
                    'Usage Complexity': complexity, 
                    'Avg Tools / Report': avg_tools_per_month, 
                    'Usage Trend': trend, 
                    'Appearances': user_data['Report Refresh Date'].nunique(),
                    'First Appearance': first_activity
                })

            self.utilized_metrics_df = pd.DataFrame(user_metrics)
            
            if not self.utilized_metrics_df.empty:
                max_consistency = self.utilized_metrics_df['Usage Consistency (%)'].max()
                max_complexity = self.utilized_metrics_df['Usage Complexity'].max()
                max_avg_complexity = self.utilized_metrics_df['Avg Tools / Report'].max()

                self.utilized_metrics_df['consistency_norm'] = self.utilized_metrics_df['Usage Consistency (%)'] / max_consistency if max_consistency > 0 else 0
                self.utilized_metrics_df['complexity_norm'] = self.utilized_metrics_df['Usage Complexity'] / max_complexity if max_complexity > 0 else 0
                self.utilized_metrics_df['avg_complexity_norm'] = self.utilized_metrics_df['Avg Tools / Report'] / max_avg_complexity if max_avg_complexity > 0 else 0
                
                self.utilized_metrics_df['Engagement Score'] = self.utilized_metrics_df['consistency_norm'] + self.utilized_metrics_df['complexity_norm'] + self.utilized_metrics_df['avg_complexity_norm']

            self.log("3. Classifying users...")
            
            # Use the latest report date as the reference for all inactivity checks
            reference_date = usage_df['Report Refresh Date'].max()
            self.log(f"Using {reference_date.strftime('%Y-%m-%d')} as the reference date for inactivity calculations.")
            
            thirty_days_ago = reference_date - timedelta(days=30)
            sixty_days_ago = reference_date - timedelta(days=60)
            ninety_days_ago = reference_date - timedelta(days=90)

            reallocation_emails, under_utilized_emails = set(), set()

            for _, row in self.utilized_metrics_df.iterrows():
                email = row['Email']
                is_new_user = pd.notna(row['First Appearance']) and row['First Appearance'] > thirty_days_ago
                
                if is_new_user:
                    under_utilized_emails.add(email)
                    continue 

                if (row['Usage Complexity'] == 0) or \
                   (pd.notna(row['Overall Recency']) and row['Overall Recency'] < ninety_days_ago) or \
                   ((row['Usage Consistency (%)'] < 25) and not is_new_user):
                    reallocation_emails.add(email)
                
                if (pd.notna(row['Overall Recency']) and (ninety_days_ago <= row['Overall Recency'] < sixty_days_ago)) or \
                   (row['Usage Trend'] == 'Decreasing') or \
                   (row['Appearances'] == 1) or \
                   ((row['Usage Consistency (%)'] < 50) and not is_new_user):
                    under_utilized_emails.add(email)

            reallocation_df = self.utilized_metrics_df[self.utilized_metrics_df['Email'].isin(reallocation_emails)].copy()
            reallocation_df['Classification'] = 'For Reallocation'
            
            under_utilized_df = self.utilized_metrics_df[self.utilized_metrics_df['Email'].isin(under_utilized_emails) & ~self.utilized_metrics_df['Email'].isin(reallocation_emails)].copy()
            under_utilized_df['Classification'] = 'Under-Utilized'

            top_utilizer_emails = set(self.utilized_metrics_df['Email']) - reallocation_emails - under_utilized_emails
            top_utilizers_df = self.utilized_metrics_df[self.utilized_metrics_df['Email'].isin(top_utilizer_emails)].copy()
            top_utilizers_df['Classification'] = 'Top Utilizer'
            
            def get_justification(row, total_months):
                reasons = []
                is_new_user = pd.notna(row['First Appearance']) and row['First Appearance'] > thirty_days_ago
                if is_new_user: reasons.append("New user (in 30-day grace period)")

                if row['Usage Complexity'] == 0: reasons.append("No tool usage recorded")
                elif pd.notna(row['Overall Recency']) and row['Overall Recency'] < ninety_days_ago: reasons.append("No activity in 90+ days")
                elif pd.notna(row['Overall Recency']) and (ninety_days_ago <= row['Overall Recency'] < sixty_days_ago): reasons.append("No activity in 60-89 days")
                
                if row['Usage Trend'] == 'Decreasing': reasons.append("Downward usage trend")
                
                active_months = int(row['Usage Consistency (%)'] * total_months / 100)
                if row['Appearances'] == 1 and not is_new_user: reasons.append("Single report appearance")
                elif row['Usage Consistency (%)'] < 50 and not is_new_user: reasons.append(f"Low consistency (active in {active_months} of {total_months} months)")
                
                return "; ".join(reasons) if reasons else "High Engagement"

            top_utilizers_df['Justification'] = "High Engagement"
            under_utilized_df['Justification'] = [get_justification(row, total_months_in_period) for _, row in under_utilized_df.iterrows()]
            reallocation_df['Justification'] = [get_justification(row, total_months_in_period) for _, row in reallocation_df.iterrows()]
            
            top_utilizers_df.sort_values(by=['Engagement Score', 'Overall Recency'], ascending=[False, False], inplace=True)
            under_utilized_df.sort_values(by=['Engagement Score', 'Overall Recency'], ascending=[True, True], inplace=True)
            reallocation_df.sort_values(by=['Engagement Score', 'Overall Recency'], ascending=[True, True], inplace=True)

            self.log("Analysis complete. Generating reports...")
            today_str = datetime.now().strftime("%d-%B-%Y")
            if not os.path.exists(self.output_folder_path): os.makedirs(self.output_folder_path)

            charts = self.create_visualizations(self.utilized_metrics_df, top_utilizers_df, under_utilized_df, matched_users_df, self.output_folder_path)
            excel_filename = os.path.join(self.output_folder_path, f"{today_str} Copilot License Evaluation.xlsx")
            self.create_excel_report(excel_filename, top_utilizers_df, under_utilized_df, reallocation_df, charts)
            word_filename = os.path.join(self.output_folder_path, f"{today_str} Copilot License Evaluation.docx")
            self.create_word_report(word_filename, top_utilizers_df, under_utilized_df, reallocation_df, charts)
            
            self.leaderboard_html_path = os.path.join(self.output_folder_path, "leaderboard.html")
            self.create_leaderboard_html(self.leaderboard_html_path, self.utilized_metrics_df)

            self.log(f"--- Success! ---")
            messagebox.showinfo("Success", f"Analysis complete. Reports saved in '{os.path.basename(self.output_folder_path)}'.")
            self.master.after(0, lambda: self.open_folder_button.pack(side="left", padx=10))
            self.master.after(0, self.activate_leaderboard_buttons)

        except PermissionError as pe: messagebox.showerror("Permission Denied", f"Could not save file. Ensure '{os.path.basename(pe.filename)}' is closed.")
        except Exception as e:
            self.log(f"An unexpected error occurred: {e}")
            import traceback
            self.log(traceback.format_exc())
            messagebox.showerror("Error", f"An unexpected error occurred: \n{e}")
        finally:
            self.master.after(0, self.reset_button)

    def reset_button(self):
        self.run_button.config(state="normal", text="Run Analysis")

    def create_visualizations(self, utilized_df, top_df, under_df, matched_df, output_folder):
        charts = {}
        plt.style.use('seaborn-v0_8-whitegrid')
        
        # Engagement Score Distribution
        if 'Engagement Score' in utilized_df.columns and not utilized_df.empty:
            plt.figure(figsize=(10, 6))
            utilized_df['Engagement Score'].plot(kind='hist', bins=20, title='Distribution of User Engagement Score')
            plt.xlabel('Engagement Score (Consistency + Complexity + Avg Tools/Rpt)')
            plt.ylabel('Number of Users')
            charts['engagement_score_hist'] = os.path.join(output_folder, 'engagement_score_hist.png')
            plt.savefig(charts['engagement_score_hist'])
            plt.close()
        
        # Tool Usage by Top Utilizers
        tool_cols = [col for col in matched_df.columns if 'Last activity date of' in col]
        top_user_activity = matched_df[matched_df['User Principal Name'].isin(top_df['Email'])]
        if not top_user_activity.empty:
            tool_usage_counts = top_user_activity[tool_cols].notna().sum().sort_values(ascending=False)
            tool_usage_counts.index = tool_usage_counts.index.str.replace('Last activity date of ', '').str.replace(' \(UTC\)', '')
            plt.figure(figsize=(12, 7))
            tool_usage_counts.plot(kind='bar', title='Most Commonly Used Tools by Top Utilizers')
            plt.ylabel('Number of Top Users Using Tool')
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            charts['top_utilizer_tools'] = os.path.join(output_folder, 'top_utilizer_tools.png')
            plt.savefig(charts['top_utilizer_tools'])
            plt.close()
        
        # Average Engagement Score Over Time for all targeted employees
        plot_data = pd.merge(
            matched_df,
            utilized_df[['Email', 'Engagement Score']],
            left_on='User Principal Name',
            right_on='Email',
            how='left'
        )
        if not plot_data.empty and 'Engagement Score' in plot_data.columns:
            trend_data = plot_data.groupby(pd.to_datetime(plot_data['Report Refresh Date']))['Engagement Score'].mean()
            
            plt.figure(figsize=(12, 6))
            trend_data.plot(kind='line', marker='o', linestyle='-', title='Average Engagement Score Over Time')
            plt.ylabel('Average Engagement Score')
            plt.xlabel('Report Date')
            plt.grid(True)
            plt.tight_layout()
            charts['avg_engagement_trend'] = os.path.join(output_folder, 'avg_engagement_trend.png')
            plt.savefig(charts['avg_engagement_trend'])
            plt.close()
            
        self.log("Visualizations created.")
        return charts

    def style_excel_sheet(self, worksheet, df):
        """Applies styling to an Excel worksheet."""
        # If the dataframe is empty, there is nothing to style.
        if df.empty:
            return

        header_fill = PatternFill(start_color="2d3748", end_color="2d3748", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        for col_num, column_title in enumerate(df.columns, 1):
            cell = worksheet.cell(row=1, column=col_num)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')

        # Zebra striping
        stripe_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        for row_index in range(2, len(df) + 2):
            if row_index % 2 == 1: # Apply to odd data rows (3, 5, 7...)
                for col_index in range(1, len(df.columns) + 1):
                    worksheet.cell(row=row_index, column=col_index).fill = stripe_fill

        for col_num, column_title in enumerate(df.columns, 1):
            max_length = 0
            column_letter = get_column_letter(col_num)
            if len(str(column_title)) > max_length:
                max_length = len(str(column_title))
            for i, cell_value in enumerate(df[column_title], 2):
                if len(str(cell_value)) > max_length:
                    max_length = len(str(cell_value))
            adjusted_width = (max_length + 2)
            worksheet.column_dimensions[column_letter].width = adjusted_width
        
        red_color, yellow_color, green_color = "F8696B", "FFEB84", "63BE7B"
        
        if 'Engagement Score' in df.columns:
            score_col_letter = get_column_letter(df.columns.get_loc('Engagement Score') + 1)
            score_range = f"{score_col_letter}2:{score_col_letter}{len(df)+1}"
            worksheet.conditional_formatting.add(score_range,
                ColorScaleRule(start_type='min', start_color=red_color,
                               mid_type='percentile', mid_value=50, mid_color=yellow_color,
                               end_type='max', end_color=green_color))

        if 'Usage Consistency (%)' in df.columns:
            consistency_col_letter = get_column_letter(df.columns.get_loc('Usage Consistency (%)') + 1)
            consistency_range = f"{consistency_col_letter}2:{consistency_col_letter}{len(df)+1}"
            worksheet.conditional_formatting.add(consistency_range,
                DataBarRule(start_type='min', end_type='max', color=green_color))

    def create_excel_report(self, filename, top_df, under_df, realloc_df, charts):
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            cols = ['Email', 'Classification', 'Usage Consistency (%)', 'Overall Recency', 'Usage Complexity', 'Avg Tools / Report', 'Usage Trend', 'Engagement Score', 'Justification']
            
            leaderboard_data = pd.concat([top_df, under_df, realloc_df]).sort_values(by="Engagement Score", ascending=False)
            
            sheets_to_process = {
                'Leaderboard': leaderboard_data,
                'Top Utilizers': top_df,
                'Under-Utilized': under_df,
                'For Reallocation': realloc_df
            }
            
            for sheet_name, df in sheets_to_process.items():
                df_to_write = df[cols].copy()
                df_to_write.insert(0, 'Rank', range(1, 1 + len(df_to_write)))
                df_to_write.to_excel(writer, sheet_name=sheet_name, index=False, float_format="%.2f")
                self.style_excel_sheet(writer.sheets[sheet_name], df_to_write)

            summary_ws = writer.book.create_sheet('Summary & Visualizations')
            from openpyxl.drawing.image import Image as OpenpyxlImage
            if 'engagement_score_hist' in charts: summary_ws.add_image(OpenpyxlImage(charts['engagement_score_hist']), 'A1')
            if 'top_utilizer_tools' in charts: summary_ws.add_image(OpenpyxlImage(charts['top_utilizer_tools']), 'A30')
            if 'avg_engagement_trend' in charts: summary_ws.add_image(OpenpyxlImage(charts['avg_engagement_trend']), 'L1')
        self.log(f"Excel report created: {filename}")
        
    def create_word_report(self, filename, top_df, under_df, realloc_df, charts):
        doc = Document()
        doc.add_heading('Copilot License Evaluation Executive Summary', level=1)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d %B %Y')}")
        doc.add_heading('User Classification Overview', level=2)
        doc.add_paragraph(f"- Top Utilizers: {len(top_df)} users")
        doc.add_paragraph(f"- Under-Utilized Users (for follow-up): {len(under_df)} users")
        doc.add_paragraph(f"- For Reallocation: {len(realloc_df)} users")
        doc.add_heading('Key Insights from Top Utilizers', level=2)
        doc.add_paragraph("Top Utilizers show high engagement. Common tools are shown below.")
        if 'top_utilizer_tools' in charts: doc.add_picture(charts['top_utilizer_tools'], width=Inches(6))
        doc.add_heading('License Reallocation Recommendations', level=2)
        doc.add_paragraph(f"{len(realloc_df)} users identified for reallocation based on inactivity or very low consistency.")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Email'
        hdr_cells[1].text = 'Consistency (%)'
        hdr_cells[2].text = 'Last Activity'
        hdr_cells[3].text = 'Justification'
        for _, row in realloc_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row['Email']
            row_cells[1].text = f"{row['Usage Consistency (%)']:.1f}"
            row_cells[2].text = row['Overall Recency'].strftime('%Y-%m-%d') if pd.notna(row['Overall Recency']) else "N/A"
            row_cells[3].text = row['Justification']
        doc.save(filename)
        self.log(f"Word report created: {filename}")

    def create_leaderboard_html(self, filename, all_users_df):
        if all_users_df is None or all_users_df.empty:
            self.log("No data available to generate leaderboard.")
            return

        leaderboard_data = all_users_df.sort_values(by="Engagement Score", ascending=False)
        
        html_head = """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Leaderboard - Haleon Theme</title>
            <script src="https://cdn.tailwindcss.com"></script>
            <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap" rel="stylesheet">
            <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
            <style>
                body { 
                    font-family: 'Inter', sans-serif; 
                    background-color: #f3f4f6;
                }
                .leaderboard-component {
                    border-radius: 1rem; 
                    box-shadow: 0 10px 15px -3px rgb(0 0 0 / 0.1), 0 4px 6px -4px rgb(0 0 0 / 0.1);
                    overflow: hidden;
                    border: 1px solid #e5e7eb;
                }
                .title-banner {
                    background-color: #000000;
                }
                .table-container {
                    background-color: #FFFFFF;
                }
                .table-header { 
                    background-color: #2d3748;
                }
                .table-row:nth-child(even) { background-color: #f9fafb; }
                .table-row:hover { 
                    background-color: #f0f0f0; 
                }
                .rank-badge { 
                    font-weight: 700; 
                    width: 2.5rem; 
                    height: 2.5rem; 
                    display: flex; 
                    align-items: center; 
                    justify-content: center; 
                    border-radius: 50%;
                    color: #000;
                }
                .progress-bar-container { 
                    background-color: #e5e7eb; 
                    border-radius: 9999px; 
                    height: 8px; 
                    width: 100%; 
                }
                .progress-bar { 
                    background: #39FF14; 
                    border-radius: 9999px; 
                    height: 100%; 
                }
                .trend-icon.Increasing { color: #16a34a; }
                .trend-icon.Stable { color: #f59e0b; }
                .trend-icon.Decreasing { color: #ef4444; }
                .trend-icon.N\\/A { color: #6b7280; }
                .neon-green-text {
                    color: #16a34a;
                }
                .user-email {
                    font-weight: 600;
                    color: #000000;
                }
            </style>
        </head>
        <body class="p-4 sm:p-6 lg:p-8">
            <div class="leaderboard-component w-full max-w-5xl mx-auto">
                <div class="title-banner p-6 text-center">
                    <h1 class="text-4xl font-bold text-white mb-2">Copilot Usage Leaderboard</h1>
                    <p class="text-gray-300">Ranking by Engagement Score</p>
                </div>
                <div class="table-container">
                    <div class="overflow-x-auto">
                        <div class="min-w-full inline-block align-middle">
                            <div class="table-header">
                                <div class="grid grid-cols-12 gap-4 px-6 py-4 text-left text-xs font-bold uppercase tracking-wider text-white">
                                    <div class="col-span-1">Rank</div>
                                    <div class="col-span-5">User</div>
                                    <div class="col-span-2 text-center">Consistency</div>
                                    <div class="col-span-2 text-center">Trend</div>
                                    <div class="col-span-2 text-right">Engagement</div>
                                </div>
                            </div>
                            <div class="divide-y divide-gray-200">
        """
        
        html_rows = ""
        max_score = leaderboard_data['Engagement Score'].max() if not leaderboard_data.empty else 3.0
        
        for rank, (_, user_row) in enumerate(leaderboard_data.iterrows(), 1):
            try:
                if pd.isna(user_row['Email']) or not isinstance(user_row['Email'], str): continue

                score_percentage = (user_row['Engagement Score'] / max_score) * 100 if max_score > 0 else 0
                hue = score_percentage * 1.2 # 0% -> 0 (red), 100% -> 120 (green)
                badge_color = f"hsl({hue}, 80%, 50%)"
                
                trend = user_row['Usage Trend']
                trend_icon_map = {'Increasing': 'fa-arrow-trend-up', 'Decreasing': 'fa-arrow-trend-down', 'Stable': 'fa-minus', 'N/A': 'fa-question'}
                trend_icon = f"fa-solid {trend_icon_map.get(trend, 'fa-minus')}"
                
                html_rows += f"""
                                <div class="grid grid-cols-12 gap-4 px-6 py-3 items-center table-row text-gray-800">
                                    <div class="col-span-1"><div class="rank-badge" style="background-color: {badge_color};"><span>{rank}</span></div></div>
                                    <div class="col-span-5"><div class="user-email">{user_row['Email']}</div></div>
                                    <div class="col-span-2 text-center"><div class="text-sm font-semibold neon-green-text">{user_row['Usage Consistency (%)']:.1f}%</div><div class="progress-bar-container mt-1"><div class="progress-bar" style="width: {user_row['Usage Consistency (%)']}%"></div></div></div>
                                    <div class="col-span-2 text-center"><i class="trend-icon {trend} {trend_icon} fa-lg"></i></div>
                                    <div class="col-span-2 text-right"><div class="text-sm font-bold neon-green-text">{user_row['Engagement Score']:.2f}</div></div>
                                </div>
                """
            except Exception as e:
                self.log(f"Error processing leaderboard row for {user_row.get('Email', 'N/A')}: {e}")
        
        html_foot = """
                            </div>
                        </div>
                    </div>
                </div>
            </body>
            </html>
        """
        
        full_html = html_head + html_rows + html_foot
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(full_html)
        self.log(f"HTML Leaderboard created: {filename}")

    def activate_leaderboard_buttons(self):
        self.open_leaderboard_button.config(state='normal')
        self.copy_html_button.config(state='normal')

    def open_leaderboard_file(self):
        if os.path.exists(self.leaderboard_html_path):
            webbrowser.open_new_tab(f"file://{os.path.realpath(self.leaderboard_html_path)}")
        else:
            messagebox.showerror("Error", "Leaderboard file not found.")
            
    def copy_leaderboard_html(self):
        if os.path.exists(self.leaderboard_html_path):
            with open(self.leaderboard_html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            self.master.clipboard_clear()
            self.master.clipboard_append(html_content)
            messagebox.showinfo("Copied!", "Leaderboard HTML source copied to clipboard.")
        else:
            messagebox.showerror("Error", "Leaderboard file not found.")

    def perform_deep_dive(self):
        user_email = self.user_email_entry.get().strip().lower()
        if not user_email: messagebox.showwarning("Input Required", "Please enter a user email."); return
        if self.full_usage_data is None or self.utilized_metrics_df is None:
            messagebox.showerror("Error", "No data loaded."); return

        user_data = self.full_usage_data[self.full_usage_data['User Principal Name'] == user_email].copy()
        user_metrics = self.utilized_metrics_df[self.utilized_metrics_df['Email'] == user_email]
        
        self.deep_dive_results.config(state='normal'); self.deep_dive_results.delete(1.0, tk.END); self.ax.clear()

        if user_data.empty or user_metrics.empty:
            self.deep_dive_results.insert(tk.END, f"No records found for '{user_email}'.")
            self.ax.set_title(f"No Data for {user_email}"); self.canvas.draw()
        else:
            metrics = user_metrics.iloc[0]
            self.deep_dive_results.insert(tk.END, f"--- Summary for {user_email} ---\n")
            self.deep_dive_results.insert(tk.END, f"Usage Consistency: {metrics['Usage Consistency (%)']:.1f}% (active in {int(metrics['Usage Consistency (%)']*len(self.full_usage_data['Report Refresh Date'].unique())/100)} of {len(self.full_usage_data['Report Refresh Date'].unique())} months)\n")
            self.deep_dive_results.insert(tk.END, f"Overall Last Activity Date: {metrics['Overall Recency'].strftime('%Y-%m-%d') if pd.notna(metrics['Overall Recency']) else 'N/A'}\n")
            self.deep_dive_results.insert(tk.END, f"Usage Complexity (Total Tools Used): {int(metrics['Usage Complexity'])}\n")
            self.deep_dive_results.insert(tk.END, f"Avg Tools / Report: {metrics['Avg Tools / Report']:.2f}\n")
            self.deep_dive_results.insert(tk.END, f"Engagement Score: {metrics['Engagement Score']:.2f}\n")
            self.deep_dive_results.insert(tk.END, f"Usage Trend: {metrics['Usage Trend']}\n\n")

            if metrics['Usage Complexity'] > 0:
                self.deep_dive_results.insert(tk.END, f"--- Detailed Records ({len(user_data)} found) ---\n")
                user_data_sorted = user_data.sort_values(by="Report Refresh Date", ascending=False)
                tool_cols = [col for col in user_data.columns if 'Last activity date of' in col]
                for _, row in user_data_sorted.iterrows():
                    report_date_str = row['Report Refresh Date'].strftime('%Y-%m-%d')
                    self.deep_dive_results.insert(tk.END, f"\nReport Date: {report_date_str}\n")
                    for col in tool_cols:
                        if pd.notna(row[col]):
                                activity_date_str = row[col].strftime('%Y-%m-%d')
                                tool_name = col.replace('Last activity date of ', '').replace(' (UTC)', ''); self.deep_dive_results.insert(tk.END, f"  - {tool_name}: {activity_date_str}\n")
            else:
                self.deep_dive_results.insert(tk.END, "--- No specific tool usage was recorded for this user. ---\n")
                report_dates = sorted(user_data['Report Refresh Date'].unique())
                if report_dates:
                    self.deep_dive_results.insert(tk.END, "Appeared in reports on the following dates:\n")
                    for date in report_dates:
                        self.deep_dive_results.insert(tk.END, f"- {pd.to_datetime(date).strftime('%Y-%m-%d')}\n")

            
            tool_cols = [col for col in self.full_usage_data.columns if 'Last activity date of' in col]
            user_data['complexity_per_report'] = user_data[tool_cols].notna().sum(axis=1)
            
            graph_data = user_data[['Report Refresh Date', 'complexity_per_report']].rename(columns={'complexity_per_report': 'complexity'})
            graph_data = graph_data.groupby('Report Refresh Date')['complexity'].sum().reset_index()
            graph_data = graph_data.sort_values(by='Report Refresh Date')

            if len(graph_data) > 1:
                self.ax.plot(graph_data['Report Refresh Date'], graph_data['complexity'], marker='o', linestyle='-')
            elif len(graph_data) == 1:
                self.ax.plot(graph_data['Report Refresh Date'], graph_data['complexity'], marker='o')

            self.ax.set_title(f"Usage Complexity per Report for {user_email}", fontsize=10)
            self.ax.set_ylabel("Number of Tools Used", fontsize=8)
            self.ax.tick_params(axis='x', rotation=45, labelsize=8)
            self.ax.grid(True)
            self.fig.tight_layout()
            self.canvas.draw()

        self.deep_dive_results.config(state='disabled')
        
    def open_report_folder(self):
        if not os.path.exists(self.output_folder_path):
            messagebox.showerror("Error", "Output folder not found.")
            return
        if sys.platform == "win32": os.startfile(self.output_folder_path)
        elif sys.platform == "darwin": subprocess.run(["open", self.output_folder_path])
        else: subprocess.run(["xdg-open", self.output_folder_path])


if __name__ == "__main__":
    root = tk.Tk()
    app = CopilotAnalysisApp(root)
    root.mainloop()